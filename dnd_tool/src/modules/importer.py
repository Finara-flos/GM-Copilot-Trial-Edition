"""文件导入与文本分段：支持 PDF / TXT / MD / DOCX。

- PDF：PyMuPDF 逐页提取文本，自动过滤页眉页脚（跨页重复行 + 页码），
  通过 "#"、Chapter、第X章 等标记识别章节结构。
- TXT / MD：UTF-8 直接读取（含 gb18030 回退）。
- DOCX：python-docx 提取段落，保留加粗/斜体/标题层级（转为 Markdown）。
- 导入后自动分段写入 SQLite（segments 表）。

注意：PyMuPDF(fitz) 在本环境中会导致「同进程其他线程的网络调用挂死」，
因此本模块对 fitz 采用惰性导入——主程序进程不应触发 fitz 加载；
PDF 解析统一走独立子进程（src/modules/pdf_worker_cli.py）。
"""
import json
import os
import re
from collections import Counter
from pathlib import Path

import docx

from src.modules.database import Database
from src.modules.llm_client import LLMError

SUPPORTED_EXTS = {".pdf", ".txt", ".md", ".markdown", ".docx"}

# 页码/页脚页眉常见行
_PAGE_NUMBER_RE = re.compile(
    r"^\s*(?:\d+\s*(?:/\s*\d+)?|第\s*[0-9一二三四五六七八九十百千]+\s*页"
    r"|page\s*\d+\s*(?:of\s*\d+)?|p\.?\s*\d+)\s*$",
    re.IGNORECASE,
)

# 章节标题模式（整块为标题时命中）
_CHAPTER_PATTERNS = [
    re.compile(r"^#{1,6}\s+.+"),                                   # Markdown 标题
    re.compile(r"^第\s*[0-9一二三四五六七八九十百千万]+\s*[章节卷部篇幕].*"),  # 中文章节/幕
    re.compile(r"^第[一二三四五六七八九十]+幕.*"),                   # 第一幕/第二幕…
    re.compile(r"^(?:chapter|section|act)\s+\d+.*", re.IGNORECASE),  # 英文章节
]

# 短行启发式标题：单行、短、无句末标点
_HEURISTIC_MIN, _HEURISTIC_MAX = 2, 40
_SENTENCE_END = re.compile(r"[。，；：、,.!?;:…]$")
_WHITESPACE = re.compile(r"\s")


def is_supported(path: str | Path) -> bool:
    """文件扩展名是否受支持。"""
    return Path(path).suffix.lower() in SUPPORTED_EXTS


def parse_document(
    path: str | Path,
    progress_cb=None,
    llm=None,
    extract_keywords: bool = True,
) -> tuple[list[dict], str, list[dict]]:
    """纯解析（不写数据库）：返回 (segments, 文件名, keywords)。

    - PDF：fast 模式本地启发式解析；smart 模式（llm 非空）逐页 LLM 清洗并抽取词库。
    - TXT / MD / DOCX：本地解析（文本类无需 LLM 精修）。
    """
    path = Path(path)
    if not path.exists():
        raise FileNotFoundError(f"文件不存在：{path}")
    ext = path.suffix.lower()
    if ext not in SUPPORTED_EXTS:
        raise ValueError(f"不支持的文件类型：{ext}（支持 {sorted(SUPPORTED_EXTS)}）")

    def progress(p: int, s: str) -> None:
        if progress_cb:
            progress_cb(max(0, min(100, p)), s)

    keywords: list[dict] = []
    if ext == ".pdf":
        if llm is not None:
            segments, filename, keywords = _parse_pdf_smart(
                path, llm, progress, extract_keywords)
        else:
            progress(5, "解析 PDF 文本…")
            text = _extract_pdf(path, progress)
            text = _separate_stat_blocks(text)
            progress(80, "正在分段…")
            segments = _segment_text(text)
            filename = path.name
            if not segments:
                segments = [{"chapter": "未分章", "content": text.strip() or "(空文档)"}]
    elif ext == ".docx":
        progress(10, "解析 DOCX 文档…")
        text = _extract_docx(path)
        progress(80, "正在分段…")
        segments = _segment_text(text)
        filename = path.name
        if not segments:
            segments = [{"chapter": "未分章", "content": text.strip() or "(空文档)"}]
    else:
        progress(10, "读取文本文件…")
        text = _read_text(path)
        progress(80, "正在分段…")
        segments = _segment_text(text)
        filename = path.name
        if not segments:
            segments = [{"chapter": "未分章", "content": text.strip() or "(空文档)"}]
    progress(100, "解析完成")
    return segments, filename, keywords


def import_document(path: str | Path, progress_cb=None, db: Database | None = None) -> tuple[list[dict], str]:
    """导入文档：解析 -> 分段 -> 入库（进程内，供非 PDF 或测试使用）。"""
    segments, filename, _ = parse_document(path, progress_cb)
    database = db or Database()
    database.clear_file(filename)
    database.insert_segments(segments, filename)
    return segments, filename


# ---------------- 智能导入（LLM 精修） ----------------
_SYSTEM_CLEAN_PROMPT = (
    "你是 DND 跑团模组文档整理助手。我会给你一页从 PDF 提取的原始文本，"
    "它可能包含：双栏排版导致的乱序、页眉页脚、页码、边栏广告或无关说明框、"
    "因换行断裂的句子。请输出：\n"
    "1. 按正确阅读顺序重排正文；\n"
    "2. 删除页眉页脚、页码、广告等与模组正文无关的内容；\n"
    "3. 章节/小节标题用 Markdown 标题（# / ##）标记；\n"
    "4. 合并被断行的句子，修正明显错误的分段：一个自然段落必须完整输出为一段，"
    "段落内部不要插入换行或空行；\n"
    "5. 只输出整理后的 Markdown 正文，不要添加原文没有的内容，不要翻译，不要解释；\n"
    "6. 怪物数据块（stat block：含 AC/HP/速度/CR/豁免 等数值行，或 TACTICS/战术 小节）"
    "必须与剧情正文分开：该页首次出现怪物数据时，先单独输出一行 '## 怪物数据'"
    "（每页最多一行），随后每个怪物单独成一段，段首用 '**怪物名**' 加粗，"
    "完整保留该怪物的数据与战术（不要删除、不要改写数值、不要翻译数值）。"
)

_KEYWORD_EXTRACT_PROMPT = (
    "从下面的模组文本中提取关键实体，按三类输出 JSON 数组：\n"
    '[{"name": "实体名", "kind": "npc|place|item", "detail": "一句话说明（可选）"}]\n'
    "npc=人物/NPC，place=地名/地点，item=关键物品/道具。\n"
    "只输出 JSON 数组，不要其他文字或解释。\n\n"
)

_LLM_CHUNK_MAX_CHARS = 6000

# 智能清洗：并发数（逐页清洗以保质量，并发提升速度）
_LLM_CONCURRENCY = 3

# 怪物数据块启发式模式
_STAT_BLOCK_LINE_RE = re.compile(
    r"(?:\bAC\s*\d+|HP\s*\d+|hp\s*\d+|速度[：: ]?\s*\d+|CR\s*[\d/]+|XP\s*[\d,]+"
    r"|豁免|Fort\s*[+-]|Ref\s*[+-]|Will\s*[+-]|Init\s*[+-]|技能[：:]|TACTICS|战术|"
    r"^(?:During Combat|Morale|战斗开始|士气))",
    re.IGNORECASE,
)
_STAT_BLOCK_HEAD_RE = re.compile(r"(?:^|\s)(?:CR|XP)\s*\d", re.IGNORECASE)


def _looks_like_stat_block(block: str) -> bool:
    """启发式判断文本块是否为怪物数据块。"""
    lines = [ln.strip() for ln in block.splitlines() if ln.strip()]
    if not lines:
        return False
    if _STAT_BLOCK_HEAD_RE.search(lines[0]):
        return True
    hits = sum(1 for ln in lines if _STAT_BLOCK_LINE_RE.search(ln))
    return hits >= 2


def _separate_stat_blocks(text: str) -> str:
    """把疑似怪物数据块的段落集中到 '## 怪物数据' 小节下（快速模式）。"""
    blocks = re.split(r"\n\s*\n", text)
    out: list[str] = []
    section_open = False
    for block in blocks:
        block = block.strip()
        if not block:
            continue
        if _looks_like_stat_block(block):
            if not section_open:
                out.append("## 怪物数据")
                section_open = True
            out.append(block)
        else:
            section_open = False
            out.append(block)
    return "\n\n".join(out)


def _merge_broken_paragraphs(text: str) -> str:
    """合并被换行/分页截断或模型拆碎的段落。

    两类合并：
    1. 前一块不以句末标点结尾（句子未写完）→ 与后一块拼回。
    2. 相邻两块都是普通文本且很短（模型常把每句独立成段）→ 合并为自然段。
    标题 / 列表 / 怪物数据块不参与合并。
    """
    blocks = [b.strip() for b in re.split(r"\n\s*\n", text) if b.strip()]
    if len(blocks) < 2:
        return text
    merged: list[str] = []
    for block in blocks:
        if merged and _should_merge_paragraph(merged[-1], block):
            merged[-1] = merged[-1] + block
        else:
            merged.append(block)
    return "\n\n".join(merged)


_SENTENCE_END_RE = re.compile(r"[。！？.!?…\"”』」）】]\s*$")

# 短段合并阈值：相邻两块都短且合并后仍较短 → 视为同一自然段被拆碎
_FRAGMENT_MAX = 90


def _should_merge_paragraph(prev: str, curr: str) -> bool:
    """判断 prev 是否应与 curr 合并为同一段。"""
    # 标题、列表、怪物数据块都不参与合并
    if _match_explicit_heading(prev) or _match_explicit_heading(curr):
        return False
    if _looks_like_stat_block(prev) or _looks_like_stat_block(curr):
        return False
    if re.match(r"^\s*(?:[-*+]|\d+[.、]|\|)", curr) \
            or re.match(r"^\s*(?:[-*+]|\d+[.、]|\|)", prev):
        return False
    # 怪物名加粗开头（如 **地精**）不合并
    if re.match(r"^\s*\*\*.+\*\*\s*$", prev) or re.match(r"^\s*\*\*.+\*\*\s*$", curr):
        return False
    # 情况 1：前一段以句末标点结尾但很短 → 模型可能把一句拆成多段，尝试合并
    if _SENTENCE_END_RE.search(prev):
        return len(prev) <= _FRAGMENT_MAX and len(prev) + len(curr) <= _FRAGMENT_MAX * 2
    # 情况 2：前一段句子未写完 → 必须拼回
    return True


def _parse_pdf_smart(
    path: Path, llm, progress, extract_keywords: bool
) -> tuple[list[dict], str, list[dict]]:
    """智能解析 PDF（不写库）：逐页 LLM 清洗 -> 分段 -> 抽取词库。"""
    # 1. 逐页粗提取（保留乱序/页眉等，交给 LLM 清洗）
    progress(5, "提取 PDF 文本…")
    pages = _extract_pdf_pages_raw(path)
    total = len(pages)
    if not total:
        raise ValueError("未能从 PDF 提取到任何文本")

    # 2. 分批清洗
    cleaned_parts: list[str] = []
    batch_index = 0
    total_batches = sum(max(1, len(_split_text(p, _LLM_CHUNK_MAX_CHARS))) for p in pages)
    for page_index, page_text in enumerate(pages):
        chunks = _split_text(page_text, _LLM_CHUNK_MAX_CHARS)
        for chunk in chunks:
            batch_index += 1
            progress(10 + int(60 * batch_index / max(total_batches, 1)),
                     f"AI 清洗批次 {batch_index}/{total_batches}…")
            cleaned_parts.append(_clean_chunk_with_llm(llm, chunk, page_index + 1))
    full_text = "\n\n".join(part for part in cleaned_parts if part.strip())

    # 3. 分段（LLM 已输出 # 标记，只认显式标题；跨页断句由分段器按章节合并）
    progress(75, "正在分段…")
    segments = _segment_text(full_text, markdown_headings=True)
    if not segments:
        segments = [{"chapter": "未分章", "content": full_text.strip() or "(空文档)"}]

    # 4. 自动抽取词库（失败不影响导入）
    keywords: list[dict] = []
    if extract_keywords:
        progress(90, "AI 抽取词库…")
        try:
            keywords = _extract_keywords_with_llm(llm, full_text)
            progress(97, f"词库抽取 {len(keywords)} 条")
        except Exception:  # noqa: BLE001  词库抽取失败不阻断导入
            progress(95, "词库抽取失败（可稍后手动添加）")

    progress(100, "解析完成")
    return segments, path.name, keywords


# ---------------- 异步智能清洗（并发 + 批量，用于子进程精修） ----------------
async def _clean_chunk_with_llm_async(client, chunk: str, page_no: int = 0) -> str:
    """单页异步清洗：发送给 LLM 并返回整理后的 Markdown。

    使用更短的读取超时（90s）与 2 次重试，避免失败时长时间挂起。
    """
    import httpx

    if page_no:
        user = f"【模组 PDF 第 {page_no} 页文本】\n---\n{chunk}\n---"
    else:
        user = f"【模组 PDF 文本】\n---\n{chunk}\n---"
    result = await client.chat(
        [{"role": "system", "content": _SYSTEM_CLEAN_PROMPT},
         {"role": "user", "content": user}],
        retries=2,
        timeout=httpx.Timeout(30.0, read=90.0),
    )
    return _strip_code_fence(result.strip())


async def _parse_pdf_smart_async(
    client,
    path: Path,
    progress,
    extract_keywords: bool,
) -> tuple[list[dict], str, list[dict]]:
    """异步智能解析 PDF：逐页清洗 + 并发（恢复单页质量，保留并发提速）。

    在子进程（pdf_worker_cli）中调用；client 需提供 async def chat。
    多页合并会降低识别质量（模型易丢页/截断），故按页单独清洗。
    """
    import asyncio

    progress(5, "提取 PDF 文本…")
    pages = _extract_pdf_pages_raw(path)
    total = len(pages)
    if not total:
        raise ValueError("未能从 PDF 提取到任何文本")

    # 逐页切块（与快速模式一致：单页上限 _LLM_CHUNK_MAX_CHARS）
    chunk_meta: list[tuple[int, str]] = []  # (page_no, chunk_text)
    for page_no, page_text in enumerate(pages, start=1):
        for chunk in _split_text(page_text, _LLM_CHUNK_MAX_CHARS):
            chunk_meta.append((page_no, chunk))
    total_chunks = len(chunk_meta)
    semaphore = asyncio.Semaphore(_LLM_CONCURRENCY)
    done = [0]

    async def clean_one(page_no: int, chunk: str) -> str:
        async with semaphore:
            result = await _clean_chunk_with_llm_async(client, chunk, page_no)
            done[0] += 1
            progress(10 + int(60 * done[0] / max(total_chunks, 1)),
                     f"AI 清洗第 {page_no} 页（{done[0]}/{total_chunks}，并发 {_LLM_CONCURRENCY}）…")
            return result

    cleaned = await asyncio.gather(
        *(clean_one(page_no, chunk) for page_no, chunk in chunk_meta))
    full_text = "\n\n".join(part for part in cleaned if part.strip())

    # 2. 分段（LLM 已输出 # 标记，只认显式标题；跨页断句由分段器按章节合并）
    progress(75, "正在分段…")
    segments = _segment_text(full_text, markdown_headings=True)
    if not segments:
        segments = [{"chapter": "未分章", "content": full_text.strip() or "(空文档)"}]

    # 3. 自动抽取词库（失败不影响导入）
    keywords: list[dict] = []
    if extract_keywords:
        progress(90, "AI 抽取词库…")
        try:
            keywords = await _extract_keywords_with_llm_async(client, full_text)
            progress(97, f"词库抽取 {len(keywords)} 条")
        except Exception:  # noqa: BLE001  词库抽取失败不阻断导入
            progress(95, "词库抽取失败（可稍后手动添加）")

    progress(100, "解析完成")
    return segments, os.path.basename(path), keywords


async def _extract_keywords_with_llm_async(client, text: str) -> list[dict]:
    """异步词库抽取。"""
    from src.modules.llm_client import LLMError

    sample = text[:8000]
    user = _KEYWORD_EXTRACT_PROMPT + "【模组文本】\n" + sample
    result = _strip_code_fence(await client.chat(
        [{"role": "user", "content": user}], retries=2, temperature=0.0))
    match = re.search(r"\[.*\]", result, re.S)
    if not match:
        return []
    try:
        data = json.loads(match.group(0))
    except json.JSONDecodeError:
        return []
    keywords: list[dict] = []
    for item in data:
        if isinstance(item, dict) and item.get("name") and item.get("kind") in ("npc", "place", "item"):
            keywords.append({
                "name": str(item["name"]).strip(),
                "kind": item["kind"],
                "detail": str(item.get("detail", "")).strip(),
            })
    return keywords


def import_document_smart(
    path: str | Path,
    llm,
    progress_cb=None,
    db: Database | None = None,
    extract_keywords: bool = True,
) -> tuple[list[dict], str]:
    """智能导入：本地粗提取 -> LLM 逐批清洗 -> 分段入库 -> 自动抽取词库。

    仅对 PDF 使用 LLM 精修；文本类（TXT/MD/DOCX）直接走本地解析。
    :param llm: 提供 chat(messages, stream_cb=...) 与 ready() 的对象
    """
    path = Path(path)
    if not path.exists():
        raise FileNotFoundError(f"文件不存在：{path}")
    ext = path.suffix.lower()
    if ext not in SUPPORTED_EXTS:
        raise ValueError(f"不支持的文件类型：{ext}（支持 {sorted(SUPPORTED_EXTS)}）")

    def progress(p: int, s: str) -> None:
        if progress_cb:
            progress_cb(max(0, min(100, p)), s)

    if ext != ".pdf":
        return import_document(path, progress_cb, db)

    segments, filename, keywords = _parse_pdf_smart(path, llm, progress, extract_keywords)
    database = db or Database()
    database.clear_file(filename)
    database.insert_segments(segments, filename)
    for kw in keywords:
        database.add_keyword(kw["name"], kw["kind"], kw.get("detail", ""), filename)
    progress(100, "导入完成")
    return segments, filename


def _extract_pdf_pages_raw(path: Path) -> list[str]:
    """逐页原始文本（保留乱序/页眉等交给 LLM）；仅去掉纯页码行与记录表页。"""
    import fitz  # 惰性导入

    doc = fitz.open(str(path))
    pages: list[str] = []
    for page in doc:
        lines = []
        for ln in page.get_text("text").splitlines():
            if ln.strip() and _PAGE_NUMBER_RE.match(ln.strip()):
                continue
            lines.append(ln)
        page_text = "\n".join(lines)
        # 跳过 PFS 记录表页（整页丢弃，非模组正文）
        if _CHRONICLE_RE.search(page_text) and len(lines) >= 5:
            hits = sum(1 for ln in lines if _CHRONICLE_RE.search(ln))
            if hits >= 3:
                continue
        pages.append(page_text)
    doc.close()
    return pages


def _split_text(text: str, max_chars: int) -> list[str]:
    """按段落边界把长文本切成小块；单个超长段落按句号边界兜底切分。"""
    if len(text) <= max_chars:
        return [text] if text.strip() else []
    parts: list[str] = []
    current = ""
    for para in re.split(r"\n\s*\n", text):
        if not para.strip():
            continue
        if len(current) + len(para) + 2 > max_chars and current:
            parts.append(current.strip())
            current = para
        else:
            current = (current + "\n\n" + para) if current else para
        # 单个段落仍超长：按句子边界拆成多个子块
        while len(current) > max_chars:
            head, rest = _cut_at_sentence(current, max_chars)
            if head:
                parts.append(head.strip())
            current = rest
    if current.strip():
        parts.append(current.strip())
    return parts


_SENTENCE_BOUNDARY_RE = re.compile(r"[。！？.!?…]")
_SENTENCE_BOUNDARY_CHARS = frozenset("。！？.!?…")


def _cut_at_sentence(text: str, max_chars: int) -> tuple[str, str]:
    """在 max_chars 之前最近的句末标点处切开，返回 (前段, 余下文本)。

    若找不到句末标点则按字符位置硬切，避免无限循环。
    """
    limit = min(max_chars, len(text))
    cut = -1
    for i in range(limit - 1, -1, -1):
        if text[i] in _SENTENCE_BOUNDARY_CHARS:
            cut = i + 1
            break
    if cut <= 0:
        cut = limit
    return text[:cut], text[cut:]


def _strip_code_fence(text: str) -> str:
    """去掉 LLM 输出可能包裹的 ```markdown 围栏。"""
    if text.startswith("```"):
        lines = text.splitlines()
        if lines and lines[0].startswith("```"):
            lines = lines[1:]
        if lines and lines[-1].strip() == "```":
            lines = lines[:-1]
        return "\n".join(lines).strip()
    return text


def _clean_chunk_with_llm(llm, chunk: str, page_no: int) -> str:
    """单批清洗：发送给 LLM 并返回整理后的 Markdown。"""
    user = f"【模组 PDF 第 {page_no} 页文本】\n---\n{chunk}\n---"
    result = llm.chat([
        {"role": "system", "content": _SYSTEM_CLEAN_PROMPT},
        {"role": "user", "content": user},
    ])
    return _strip_code_fence(result.strip())


def _extract_keywords_with_llm(llm, text: str) -> list[dict]:
    """从文本中抽取 {name, kind, detail} 词条列表。"""
    sample = text[:8000]
    user = _KEYWORD_EXTRACT_PROMPT + "【模组文本】\n" + sample
    result = _strip_code_fence(llm.chat([
        {"role": "user", "content": user},
    ], temperature=0.0))
    match = re.search(r"\[.*\]", result, re.S)
    if not match:
        return []
    try:
        data = json.loads(match.group(0))
    except json.JSONDecodeError:
        return []
    keywords: list[dict] = []
    for item in data:
        if isinstance(item, dict) and item.get("name") and item.get("kind") in ("npc", "place", "item"):
            keywords.append({
                "name": str(item["name"]).strip(),
                "kind": item["kind"],
                "detail": str(item.get("detail", "")).strip(),
            })
    return keywords


# ---------------- PDF ----------------
# 字号比正文大多少视为标题
_HEADING_SIZE_RATIO = 1.15

# PFS 记录表（Chronicle Sheet）特征：含勾选框、金币价格、记录表栏目
_CHRONICLE_RE = re.compile(
    r"(?:☐|□|Chronicle|Pathfinder Society #|TOTAL VALUE OF ITEMS SOLD|"
    r"Items Sold|GP Gained|GM's Hat|GM’s Hat|Prestige Award|Faction\b|"
    r"\d[\d,]*\s*gp\)|Tier 5–6|Tier 8–9)",
    re.IGNORECASE,
)

# 装饰线 / 分隔符行
_RULE_LINE_RE = re.compile(r"^[─━\-═—=]{6,}$")

# 翻译附加信息（首页注记：QQ 群 / 论坛链接等非模组内容）
_BOILERPLATE_RE = re.compile(
    r"(?:QQ 讨论群|PFS 果园|http://|https://|bbs/|翻译[：:]|译者[：:])"
)


def _is_chronicle_page(lines: list[dict]) -> bool:
    """判断整页是否为 PFS 记录表（Chronicle Sheet），是则整页丢弃。

    记录表是 PFS 模组最后一页的玩家结算表：勾选框、金币价格、
    "Items Sold" 等栏目，不属于模组正文。
    """
    if len(lines) < 5:
        return False
    hits = sum(1 for ln in lines if _CHRONICLE_RE.search(ln["text"]))
    return hits >= 3


def _extract_pdf(path: Path, progress) -> str:
    """逐页提取 PDF 文本行（含字号/坐标），过滤页眉页脚与记录表，
    按阅读顺序排序（双栏识别），字号识别标题。"""
    import fitz  # 惰性导入：仅在子进程/主线程 PDF 场景加载

    doc = fitz.open(str(path))
    total = len(doc)
    pages: list[list[dict]] = []
    page_w, page_h = 595.0, 842.0
    for index, page in enumerate(doc):
        page_w, page_h = page.rect.width, page.rect.height
        lines: list[dict] = []
        data = page.get_text("dict")
        for block in data.get("blocks", []):
            if block.get("type") != 0:  # 0 = 文本块，跳过图片块
                continue
            for line in block.get("lines", []):
                spans = line.get("spans") or []
                if not spans:
                    continue
                text = "".join(s.get("text", "") for s in spans)
                text = " ".join(text.split())  # 统一空白
                if not text:
                    continue
                x0, y0, x1, y1 = line.get("bbox", (0.0, 0.0, 0.0, 0.0))
                lines.append({
                    "x0": x0, "y0": y0, "x1": x1, "y1": y1,
                    "text": text,
                    "size": max((s.get("size") or 0) for s in spans),
                })
        # 跳过 PFS 记录表页（整页丢弃，非模组正文）
        if _is_chronicle_page(lines):
            progress(10 + int(55 * (index + 1) / max(total, 1)),
                     f"跳过记录表页（第 {index + 1} 页）")
            continue
        pages.append(lines)
        progress(10 + int(55 * (index + 1) / max(total, 1)),
                 f"解析 PDF 第 {index + 1}/{total} 页")
    doc.close()
    progress(68, "过滤页眉页脚…")
    pages = _filter_pdf_header_footer(pages, page_h)
    progress(74, "识别章节与排版…")
    return _layout_pdf_lines(pages, page_w, page_h)


def _filter_pdf_header_footer(
    pages: list[list[dict]], page_height: float
) -> list[list[dict]]:
    """按顶部/底部区域统计跨页重复行，移除页眉页脚与页码。"""
    n_pages = len(pages)
    if n_pages < 2:
        return pages
    top_limit = page_height * 0.10
    bottom_limit = page_height * 0.90
    top_counter: Counter = Counter()
    bottom_counter: Counter = Counter()
    for lines in pages:
        for ln in lines:
            if ln["y1"] < top_limit:
                top_counter[ln["text"]] += 1
            elif ln["y0"] > bottom_limit:
                bottom_counter[ln["text"]] += 1
    min_count = max(2, int(n_pages * 0.6))
    repeated_top = {t for t, c in top_counter.items() if c >= min_count}
    repeated_bottom = {t for t, c in bottom_counter.items() if c >= min_count}

    result: list[list[dict]] = []
    for lines in pages:
        kept = []
        for ln in lines:
            if ln["y1"] < top_limit and (
                ln["text"] in repeated_top or _PAGE_NUMBER_RE.match(ln["text"])
            ):
                continue
            if ln["y0"] > bottom_limit and (
                ln["text"] in repeated_bottom or _PAGE_NUMBER_RE.match(ln["text"])
            ):
                continue
            kept.append(ln)
        result.append(kept)
    return result


def _find_column_gap(lines: list[dict], text_min_x: float, text_width: float):
    """检测双栏排版，返回栏间分割 x 坐标；单栏返回 None。

    只统计"正文行"（行宽足够、非页眉页脚区域、非短标注）的 x0 分布，
    按 x0 聚类成簇，簇间空隙即栏间隙。避免页码/地图标注/页眉等
    落在栏间隙的行干扰检测。
    """
    n = len(lines)
    if n < 6:
        return None
    # 正文行筛选：行宽占文本区一定比例，避免短标注行（页码、地图名、装饰）
    body_lines = [
        ln for ln in lines
        if (ln["x1"] - ln["x0"]) >= text_width * 0.15
        and ln["x0"] - text_min_x < text_width * 0.8
    ]
    if len(body_lines) < 6:
        body_lines = lines
    ordered = sorted(body_lines, key=lambda ln: ln["x0"])
    # 簇合并阈值：栏间隙通常远大于同一栏内行首缩进的差异。
    # 用"行 x1 的平均宽度"估算栏宽，间隙阈值取栏宽的 1/6 左右；
    # 同时用固定下限（12pt）避免把同栏缩进行拆成两簇。
    avg_width = sum(ln["x1"] - ln["x0"] for ln in ordered) / len(ordered)
    merge_tol = max(12.0, avg_width * 0.12)
    clusters: list[list[dict]] = []
    for ln in ordered:
        if clusters and ln["x0"] - clusters[-1][-1]["x0"] <= merge_tol:
            clusters[-1].append(ln)
        else:
            clusters.append([ln])
    # 保留成员足够多的簇（至少 12% 的正文行），找两个最大簇之间的空隙
    big = [c for c in clusters if len(c) >= max(2, len(ordered) * 0.12)]
    if len(big) < 2:
        return None
    big.sort(key=lambda c: sum(ln["x0"] for ln in c) / len(c))
    best_gap, best_split = 0, None
    for i in range(1, len(big)):
        prev_max_x1 = max(ln["x1"] for ln in big[i - 1])
        curr_min_x0 = min(ln["x0"] for ln in big[i])
        gap = curr_min_x0 - prev_max_x1
        if gap > best_gap:
            best_gap, best_split = gap, i
    if best_split is None or best_gap < 8.0:
        return None
    left, right = big[:best_split], big[best_split:]
    left_max_x1 = max(ln["x1"] for c in left for ln in c)
    right_min_x0 = min(ln["x0"] for c in right for ln in c)
    if left_max_x1 < right_min_x0:
        return (left_max_x1 + right_min_x0) / 2
    return None


def _order_lines(lines: list[dict], page_width: float) -> list[dict]:
    """按阅读顺序排序：双栏先左后右；单栏主栏在前、窄边栏后置。"""
    if len(lines) < 4:
        result = sorted(lines, key=lambda ln: (ln["y0"], ln["x0"]))
        for ln in result:
            ln["col"] = 0
        return result

    text_min_x = min(ln["x0"] for ln in lines)
    text_max_x = max(ln["x1"] for ln in lines)
    text_width = max(text_max_x - text_min_x, 1.0)

    gap = _find_column_gap(lines, text_min_x, text_width)
    if gap is not None:
        left = sorted((ln for ln in lines if ln["x1"] <= gap),
                      key=lambda ln: (ln["y0"], ln["x0"]))
        right = sorted((ln for ln in lines if ln["x0"] >= gap),
                       key=lambda ln: (ln["y0"], ln["x0"]))
        for ln in left:
            ln["col"] = 0
        for ln in right:
            ln["col"] = 1
        return left + right

    # 单栏：主栏 + 窄边栏（宽度不足文本区 45% 且右偏 → 视为边栏/浮框，后置）
    main, marginal = [], []
    for ln in lines:
        width = ln["x1"] - ln["x0"]
        offset = ln["x0"] - text_min_x
        if width < text_width * 0.45 and offset > text_width * 0.08:
            ln["col"] = 2
            marginal.append(ln)
        else:
            ln["col"] = 0
            main.append(ln)
    main.sort(key=lambda ln: (ln["y0"], ln["x0"]))
    marginal.sort(key=lambda ln: (ln["y0"], ln["x0"]))
    return main + marginal


_CJK_PUNCT_END = re.compile(r"\s+([。，；：、？！」』）】】])")
_CJK_PUNCT_START = re.compile(r"([（「『【])\s+")
_CJK_SPACE = re.compile(r"(?<=[\u4e00-\u9fff])\s+(?=[\u4e00-\u9fff])")


def _clean_cjk_text(text: str) -> str:
    """清理中英混排提取时的多余空格。"""
    text = _CJK_PUNCT_END.sub(r"\1", text)
    text = _CJK_PUNCT_START.sub(r"\1", text)
    text = _CJK_SPACE.sub("", text)
    return text


def _group_lines_into_paragraphs(
    ordered_lines: list[dict], body_size: float
) -> str:
    """把排序后的文本行聚合成段落；标题行加 '#' 前缀并独立成段。"""
    line_height = max(body_size * 1.4, 6.0)
    paragraphs: list[str] = []
    current: list[str] = []
    last_y1 = None
    last_col = None
    for ln in ordered_lines:
        if ln.get("heading"):
            if current:
                paragraphs.append(_clean_cjk_text(" ".join(current)))
                current = []
            text = ln["text"]
            # 源文本已带 # 时不再重复加前缀
            heading_line = text if text.lstrip().startswith("#") else f"# {text}"
            paragraphs.append(heading_line)
            last_y1 = ln["y1"]
            last_col = ln.get("col")
            continue
        col = ln.get("col", 0)
        new_block = (
            last_y1 is not None and ln["y0"] - last_y1 > line_height * 1.7
        ) or (last_col is not None and col != last_col)
        if current and new_block:
            paragraphs.append(_clean_cjk_text(" ".join(current)))
            current = []
        current.append(ln["text"])
        last_y1 = ln["y1"]
        last_col = col
    if current:
        paragraphs.append(_clean_cjk_text(" ".join(current)))
    return "\n\n".join(paragraphs)


def _layout_pdf_lines(pages: list[list[dict]], page_w: float, page_h: float) -> str:
    """识别正文字号与标题，按阅读顺序输出为带空行的文本。

    过滤装饰线与翻译附加信息行；页与页之间保留空行，交由
    _merge_broken_paragraphs 合并跨页断裂的句子。
    """
    out_pages: list[str] = []
    for lines in pages:
        if not lines:
            continue
        # 过滤装饰线（────）与附加信息（QQ 群 / 链接 / 译者注）
        lines = [
            ln for ln in lines
            if not _RULE_LINE_RE.match(ln["text"])
            and not _BOILERPLATE_RE.search(ln["text"])
        ]
        if not lines:
            continue
        sizes = sorted(ln["size"] for ln in lines if ln["size"] > 0)
        body_size = sizes[len(sizes) // 2] if sizes else 10.0
        for ln in lines:
            # 字号判标题需同时满足：短行、不含任何句读（避免折行/长句误判）、
            # 非怪物数据特征行（子级别/CR/AC/HP 标签不是章节标题）
            text = ln["text"]
            size_heading = (
                ln["size"] >= body_size * _HEADING_SIZE_RATIO
                and len(text) <= 50
                and not re.search(r"[。！？.!?，；：、,;:]", text)
                and not re.search(r"子级别\s*[\d–\-—]", text)
                and not _STAT_BLOCK_LINE_RE.search(text)
            )
            is_heading = size_heading or _match_explicit_heading(text) is not None
            ln["heading"] = is_heading
        ordered = _order_lines(lines, page_w)
        out_pages.append(_group_lines_into_paragraphs(ordered, body_size))
    return "\n\n".join(page for page in out_pages if page.strip())


# ---------------- TXT / MD / DOCX ----------------
def _read_text(path: Path) -> str:
    """读取 UTF-8 文本；失败时回退 gb18030，再失败则容错读取。"""
    for encoding in ("utf-8", "utf-8-sig", "gb18030"):
        try:
            return Path(path).read_text(encoding=encoding)
        except (UnicodeDecodeError, LookupError):
            continue
    return Path(path).read_text(encoding="utf-8", errors="replace")


def _extract_docx(path: Path) -> str:
    """提取 DOCX 段落与表格，转为 Markdown（标题/加粗/斜体保留）。"""
    document = docx.Document(str(path))
    parts: list[str] = []

    for para in document.paragraphs:
        text = _runs_to_markdown(para)
        if not text.strip():
            continue
        style_name = (para.style.name or "") if para.style else ""
        if "heading" in style_name.lower():
            digits = "".join(ch for ch in style_name if ch.isdigit())
            level = int(digits) if digits else 1
            text = f"{'#' * min(level, 6)} {text}"
        parts.append(text)

    for table in document.tables:
        parts.append(_table_to_markdown(table))

    return "\n\n".join(parts)


def _runs_to_markdown(para) -> str:
    """把段落内的 run 样式转为 Markdown：加粗 **、斜体 *。"""
    out = ""
    for run in para.runs:
        text = run.text
        if not text:
            continue
        if run.bold:
            text = f"**{text}**"
        if run.italic:
            text = f"*{text}*"
        out += text
    return out or para.text


def _table_to_markdown(table) -> str:
    """Word 表格转 Markdown 表格。"""
    rows = []
    for row in table.rows:
        cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
        rows.append("| " + " | ".join(cells) + " |")
    if not rows:
        return ""
    separator = "| " + " | ".join("---" for _ in range(len(rows[0].split("|")) - 2)) + " |"
    return "\n".join([rows[0], separator] + rows[1:])


# ---------------- 分段 ----------------
def _match_explicit_heading(line: str) -> str | None:
    """匹配显式标题模式（# / Chapter / 第X章），返回清洗后的标题。"""
    line = line.strip()
    if not line:
        return None
    for pattern in _CHAPTER_PATTERNS:
        if pattern.match(line):
            return re.sub(r"^#{1,6}\s*", "", line).strip()
    return None


def _match_heading(block: str) -> str | None:
    """判断文本块是否为章节标题；是则返回清洗后的标题文本。

    显式模式（#/Chapter/第X章）可用于块内首行；
    短行启发式仅用于"整块只有一行"的情况，避免误伤列表/表格/代码块。
    """
    block = block.strip()
    lines = block.splitlines()
    first_line = lines[0].strip() if lines else ""

    explicit = _match_explicit_heading(first_line)
    if explicit:
        return explicit
    # 启发式：整块只有一行、长度适中、不含句内/句末句读（标题不应有。！？；：）
    if len(lines) == 1:
        length = len(first_line)
        # 忽略行首编号点（1. 2、 等），其余位置不应含句读
        stripped_head = re.sub(r"^\d+[.、]\s*", "", first_line)
        if (_HEURISTIC_MIN <= length <= _HEURISTIC_MAX
                and not re.search(r"[。！？.!?；：、,;:]", stripped_head)):
            if "://" not in first_line and not _PAGE_NUMBER_RE.match(first_line):
                # Markdown 分隔线（--- / === / *** / ___）不是标题
                if re.match(r"^[─━\-═—=*_~]{3,}$", first_line):
                    return None
                # 怪物数据块标题（以 子级别/CR/XP 开头，或怪物名+CR 组合）不是章节
                if _looks_like_stat_block_title(first_line):
                    return None
                return first_line
    return None


def _looks_like_stat_block_title(line: str) -> bool:
    """判断单行是否为怪物数据块标题（非章节标题）。

    命中形态：
    - 以「子级别」+ 数字/CR 开头（难度分级标签）
    - 以 CR/XP + 数字开头（挑战等级标签）
    - 行首为怪物名，行内含 CR/AC/HP 等特征（如 "亡灵士兵 (3) CR 2 …"）
    """
    line = line.strip()
    if re.match(r"^子级别\s*[\d–\-—]", line):
        return True
    if re.match(r"^(?:CR|XP)\s*\d", line, re.IGNORECASE):
        return True
    # 行首有实体名 + 行内含 CR 数字（如 亡灵士兵 (3) CR 2 人类骷髅勇士）
    if re.search(r"\bCR\s*\d", line, re.IGNORECASE) and not re.match(r"^\d+[.、]", line):
        return True
    return False


def _segment_text(text: str, markdown_headings: bool = False) -> list[dict]:
    """按空行分段，遇到标题则切换当前章节。

    :param markdown_headings: True=文本来自 LLM 精修（已带 # 标记，只认显式
        标题，避免启发式误判正文短句）；False=快速本地解析（启用启发式标题）。

    - `# 一级标题`：切换顶层章节（概述/冒险开始/第X幕/怪物数据/结团…）
    - `## 二级及以下标题`：作为当前章节内的小节标题，保留在内容中，
      不切换顶层章节（stat block 的"数据/战术/生物"等小节属此列）
    - Markdown 分隔线（---）与页码行（第 N 页）直接丢弃
    - 跨页断句合并：同一章节内，前一段不以句末标点结尾（句子在分页处
      被截断），自动与后一段拼接，恢复完整段落。
    """
    blocks = re.split(r"\n\s*\n", text)
    segments: list[dict] = []
    chapter = "未分章"
    for block in blocks:
        block = block.strip()
        if not block:
            continue
        lines = block.splitlines()
        first_line = lines[0].strip() if lines else ""
        # 丢弃分隔线与页码行
        if re.match(r"^[─━\-═—=*_~]{3,}$", first_line):
            continue
        if _PAGE_NUMBER_RE.match(first_line) or re.match(r"^第\s*\d+\s*页$", first_line):
            continue
        heading = _match_heading_mode(block, markdown_headings)
        if heading:
            level = _heading_level(block)
            # "怪物数据" 无论 # 级别都作为顶层章节（用户要求怪物数据独立成章）
            if heading.strip() == "怪物数据":
                level = 1
            if level <= 1:
                # 一级标题：切换顶层章节
                chapter = heading
                rest = "\n".join(lines[1:]).strip() if len(lines) > 1 else ""
                if rest:
                    segments.append({"chapter": chapter, "content": rest})
            else:
                # 二级及以下：小节标题 + 其下正文合并为一段，归入当前章节
                rest = "\n".join(lines[1:]).strip() if len(lines) > 1 else ""
                if rest:
                    segments.append({"chapter": chapter, "content": block})
                else:
                    # 标题单独成块：并入前一段末尾（下段正文将与之相连）
                    segments.append({"chapter": chapter, "content": first_line})
            continue
        # 跨页断句/小节标题合并：前段是未完结句子，或前段是纯二级标题行
        prev = segments[-1] if segments else None
        prev_is_subheading = bool(
            prev and prev["chapter"] == chapter
            and re.match(r"^#{2,6}\s+.+$", prev["content"].strip())
            and len(prev["content"].splitlines()) == 1
        )
        # 显式一级标题（# 概述）永不合并；其余（断句/二级标题）可合并
        prev_mergeable = bool(
            prev and prev["chapter"] == chapter
            and not _SENTENCE_END_RE.search(prev["content"])
            and not _looks_like_stat_block(prev["content"])
            and not re.match(r"^#\s+.+$", prev["content"].strip())
        )
        if prev_mergeable:
            # 纯二级标题行并入正文时用换行分隔；断句合并则直接拼接
            separator = "\n" if prev_is_subheading else ""
            segments[-1]["content"] = segments[-1]["content"] + separator + block
        else:
            segments.append({"chapter": chapter, "content": block})
    return segments


def _match_heading_mode(block: str, markdown_headings: bool) -> str | None:
    """按模式判断标题：精修模式只认显式 # 标题；快速模式加启发式。"""
    if markdown_headings:
        # 只认显式 Markdown 标题（# / ## ...）
        first_line = block.strip().splitlines()[0].strip() if block.strip() else ""
        m = re.match(r"^#{1,6}\s+(.+)$", first_line)
        if not m:
            return None
        title = re.sub(r"^#{1,6}\s*", "", first_line).strip()
        # 模型可能误把正文标成标题：页码、问句、引文（含句末标点）、纯符号
        if re.match(r"^第\s*\d+\s*页$", title):
            return None
        if "？" in title or "?" in title:
            return None
        if re.search(r"[。！!]", title):
            return None
        if re.match(r"^[─━\-═—=*_~]{3,}$", title):
            return None
        return title
    return _match_heading(block)


def _heading_level(block: str) -> int:
    """返回块首标题的 Markdown 级别（# 数量）；非标题返回 0。"""
    first_line = block.strip().splitlines()[0].strip() if block.strip() else ""
    m = re.match(r"^(#{1,6})\s+", first_line)
    return len(m.group(1)) if m else 0
